"""Shared helper utilities for Gmail MCP server."""

import base64
import sys
from html.parser import HTMLParser
from typing import Any, Optional


def format_size(num_bytes: int) -> str:
    """Convert bytes to human-readable size string."""
    if num_bytes < 1024:
        return f"{num_bytes} B"
    elif num_bytes < 1024 * 1024:
        return f"{num_bytes / 1024:.1f} KB"
    elif num_bytes < 1024 * 1024 * 1024:
        return f"{num_bytes / (1024 * 1024):.1f} MB"
    else:
        return f"{num_bytes / (1024 * 1024 * 1024):.1f} GB"


class _HTMLStripper(HTMLParser):
    """Minimal HTML-to-text converter using stdlib only."""

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_tags = {"script", "style", "head"}
        self._block_tags = {
            "p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6",
            "li", "tr", "blockquote", "pre",
        }
        self._current_skip: list[str] = []

    def handle_starttag(self, tag: str, attrs: list) -> None:
        tag_lower = tag.lower()
        if tag_lower in self._skip_tags:
            self._current_skip.append(tag_lower)
        if tag_lower in self._block_tags:
            self._parts.append("\n")
        if tag_lower == "a":
            # Preserve link text; href already shown via text in most emails
            pass

    def handle_endtag(self, tag: str) -> None:
        tag_lower = tag.lower()
        if self._current_skip and self._current_skip[-1] == tag_lower:
            self._current_skip.pop()
        if tag_lower in self._block_tags:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._current_skip:
            self._parts.append(data)

    def handle_entityref(self, name: str) -> None:
        # Common HTML entities
        entities = {
            "amp": "&", "lt": "<", "gt": ">", "quot": '"',
            "apos": "'", "nbsp": " ",
        }
        if not self._current_skip:
            self._parts.append(entities.get(name, ""))

    def handle_charref(self, name: str) -> None:
        if not self._current_skip:
            try:
                if name.startswith("x"):
                    char = chr(int(name[1:], 16))
                else:
                    char = chr(int(name))
                self._parts.append(char)
            except (ValueError, OverflowError):
                pass

    def get_text(self) -> str:
        text = "".join(self._parts)
        # Collapse excessive blank lines
        import re
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def strip_html(html: str) -> str:
    """Strip HTML tags and return plain text."""
    stripper = _HTMLStripper()
    try:
        stripper.feed(html)
        return stripper.get_text()
    except Exception:
        # Fallback: crude regex strip
        import re
        text = re.sub(r"<[^>]+>", " ", html)
        return re.sub(r"\s+", " ", text).strip()


def extract_body(payload: dict) -> str:
    """
    Recursively extract email body from a Gmail message payload.
    Prefers text/plain; falls back to text/html (stripped).
    Returns empty string if no body found.
    """
    plain: str = ""
    html: str = ""

    def _decode_data(data: str) -> str:
        try:
            padded = data + "=" * (4 - len(data) % 4)
            return base64.urlsafe_b64decode(padded).decode("utf-8", errors="replace")
        except Exception:
            return ""

    def _recurse(part: dict) -> None:
        nonlocal plain, html
        mime = part.get("mimeType", "")
        body = part.get("body", {})
        data = body.get("data", "")

        if mime == "text/plain" and data and not plain:
            plain = _decode_data(data)
        elif mime == "text/html" and data and not html:
            html = _decode_data(data)
        elif mime.startswith("multipart/"):
            for sub_part in part.get("parts", []):
                _recurse(sub_part)

    _recurse(payload)

    if plain:
        return plain
    if html:
        return strip_html(html)
    return "(No readable body found)"


def extract_attachment_info(payload: dict) -> list[dict]:
    """
    Recursively collect attachment metadata from a Gmail message payload.
    Returns list of dicts with keys: filename, mimeType, size, attachmentId.
    """
    attachments: list[dict] = []

    def _recurse(part: dict) -> None:
        filename = part.get("filename", "")
        body = part.get("body", {})
        if filename:
            size_bytes = body.get("size", 0)
            attachments.append(
                {
                    "filename": filename,
                    "mimeType": part.get("mimeType", "application/octet-stream"),
                    "size": format_size(size_bytes),
                    "size_bytes": size_bytes,
                    "attachmentId": body.get("attachmentId", ""),
                }
            )
        for sub_part in part.get("parts", []):
            _recurse(sub_part)

    _recurse(payload)
    return attachments


def _try_pdf(binary: bytes) -> tuple[Optional[str], Optional[str]]:
    from io import BytesIO
    try:
        import pypdf
        reader = pypdf.PdfReader(BytesIO(binary))
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            page_text = (page.extract_text() or "").strip()
            parts.append(f"--- Page {i + 1} ---\n{page_text or '(empty/scanned page)'}")
        return "\n\n".join(parts), f"PDF extracted via pypdf ({len(reader.pages)} page(s))"
    except Exception as e:
        print(f"[gmail-mcp] PDF parse failed: {e}", file=sys.stderr)
        return None, None


def _try_docx(binary: bytes) -> tuple[Optional[str], Optional[str]]:
    from io import BytesIO
    try:
        import docx  # python-docx
        doc = docx.Document(BytesIO(binary))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(paragraphs).strip()
        if not text:
            return None, None
        return text, "DOCX extracted via python-docx"
    except Exception as e:
        print(f"[gmail-mcp] DOCX parse failed: {e}", file=sys.stderr)
        return None, None


def _try_xlsx(binary: bytes) -> tuple[Optional[str], Optional[str]]:
    from io import BytesIO
    try:
        import openpyxl
        wb = openpyxl.load_workbook(BytesIO(binary), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                parts.append("\t".join(cells))
        text = "\n".join(parts).strip()
        if not text:
            return None, None
        return text, f"XLSX extracted via openpyxl ({len(wb.worksheets)} sheet(s))"
    except Exception as e:
        print(f"[gmail-mcp] XLSX parse failed: {e}", file=sys.stderr)
        return None, None


def _try_utf8(binary: bytes) -> tuple[Optional[str], Optional[str]]:
    try:
        text = binary.decode("utf-8")
    except UnicodeDecodeError:
        return None, None
    # Heuristic: only treat as text if mostly printable
    if not text:
        return None, None
    printable = sum(1 for c in text if c.isprintable() or c in "\n\r\t")
    if printable / len(text) < 0.85:
        return None, None
    return text, "decoded as UTF-8"


def extract_attachment_text(binary: bytes, mime_type: str, filename: str) -> tuple[Optional[str], Optional[str]]:
    """
    Try to extract human-readable text from attachment bytes.
    Returns (text, format_label) — both None if we can't extract.

    Strategy: prefer hints (filename + MIME), but ALWAYS verify by actually
    parsing the bytes. This is robust even when Gmail returns
    application/octet-stream or no filename.
    """
    mime = (mime_type or "").lower()
    name_lower = (filename or "").lower()

    # Magic-byte sniffing — most reliable
    is_pdf = binary[:5] == b"%PDF-"
    is_zip = binary[:4] == b"PK\x03\x04"

    # ----- PDF -----
    if is_pdf or mime == "application/pdf" or name_lower.endswith(".pdf"):
        result = _try_pdf(binary)
        if result[0] is not None:
            return result

    # ----- ZIP-based Office formats: try DOCX, then XLSX -----
    if is_zip or any(name_lower.endswith(ext) for ext in (".docx", ".xlsx", ".xlsm")):
        # Try DOCX first (most common)
        if "wordprocessingml" in mime or name_lower.endswith(".docx") or is_zip:
            result = _try_docx(binary)
            if result[0] is not None:
                return result
        # Then XLSX
        if "spreadsheetml" in mime or name_lower.endswith((".xlsx", ".xlsm")) or is_zip:
            result = _try_xlsx(binary)
            if result[0] is not None:
                return result

    # ----- Plain text family (decode by MIME or extension) -----
    if (
        mime.startswith("text/")
        or mime in ("application/json", "application/xml", "application/javascript",
                    "application/x-yaml", "application/yaml")
        or name_lower.endswith((".txt", ".md", ".csv", ".tsv", ".json", ".xml",
                                 ".yaml", ".yml", ".html", ".htm", ".log"))
    ):
        try:
            return binary.decode("utf-8", errors="replace"), "decoded as UTF-8"
        except Exception:
            pass

    # ----- Last resort: UTF-8 sniff (catches text/* with weird/missing MIME) -----
    result = _try_utf8(binary)
    if result[0] is not None:
        return result

    return None, None


def format_email_headers(headers: list[dict]) -> dict[str, str]:
    """
    Convert Gmail's header list format to a clean dict.
    Input:  [{"name": "Subject", "value": "Hello"}, ...]
    Output: {"Subject": "Hello", ...}
    Duplicate header names: last value wins.
    """
    result: dict[str, str] = {}
    for h in headers:
        name = h.get("name", "")
        value = h.get("value", "")
        if name:
            result[name] = value
    return result


def handle_api_error(e: Exception) -> str:
    """
    Return a user-friendly, actionable error message for Google API errors.
    Never exposes raw tracebacks.
    """
    import googleapiclient.errors  # type: ignore

    print(f"[gmail-mcp] API error: {type(e).__name__}: {e}", file=sys.stderr)

    if isinstance(e, googleapiclient.errors.HttpError):
        status = e.resp.status if hasattr(e, "resp") else 0
        if status == 400:
            return "Bad request — check your query syntax or parameters and try again."
        elif status == 401:
            return "Authentication error — the token may be expired. Use gmail_authenticate to re-authenticate this account."
        elif status == 403:
            return "Permission denied — you may not have access to this resource, or the token needs refreshing. Use gmail_authenticate to re-authenticate."
        elif status == 404:
            return "Resource not found — check the message ID or label name. Run gmail_search_emails first to get valid message IDs."
        elif status == 429:
            return "Rate limited by Gmail API — please wait a minute and try again."
        elif status == 500:
            return "Gmail server error — this is a temporary issue on Google's side. Try again in a moment."
        elif status == 503:
            return "Gmail service unavailable — check your internet connection and try again."
        else:
            return f"Gmail API error (HTTP {status}) — please try again or re-authenticate the account."
    elif isinstance(e, TimeoutError):
        return "Request timed out — check your internet connection and try again."
    elif isinstance(e, ConnectionError):
        return "Network connection error — check your internet connection and try again."
    elif isinstance(e, FileNotFoundError):
        return str(e)
    else:
        # Generic fallback — never expose traceback
        return f"An unexpected error occurred: {type(e).__name__}. Check the server logs for details."


def validate_account_alias(alias: str, config: dict) -> str | None:
    """
    Validate that an account alias exists in the config.
    Returns None if valid, or an error string if not found.
    """
    accounts = config.get("accounts", {})
    if alias not in accounts:
        available = ", ".join(sorted(accounts.keys())) if accounts else "(none configured)"
        return (
            f"Account '{alias}' not found. "
            f"Available accounts: {available}. "
            f"Use gmail_authenticate to add a new account, or gmail_list_accounts to see all accounts."
        )
    return None
